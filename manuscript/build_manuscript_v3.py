"""
Build manuscript/manuscript.v3.docx from manuscript/manuscript.v3.md.

Parses the markdown subset used in the v3 manuscript (headings, paragraphs,
bold/italic spans, numbered references, and figure entries) and renders a
Word document with a separate title page, embedded figures, and page numbers.

Usage:
    python build_manuscript_v3.py
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parent.parent
MD = ROOT / "manuscript" / "manuscript.v3.md"
OUT = ROOT / "manuscript" / "manuscript.v3.docx"

FONT = "Times New Roman"
BODY_SIZE = Pt(11)
CAPTION_SIZE = Pt(9.5)

# Cap displayed width per figure (inches). Wide/tall figures get narrower so
# the figures section stays compact.
FIGURE_WIDTHS = {
    "fig1_three_track_architecture.png": 5.4,
    "fig2_phase5_rl_results.png": 5.4,
    "fig3_biomarker_decision_rule.png": 4.6,
    "fig4_fno_speedup.png": 5.4,
    "fig5_circadian_ppo.png": 5.4,
    "fig6_virtual_trial_results.png": 5.4,
    "fig7_cross_track_radar.png": 4.0,
    "fig8_roadmap_timeline.png": 5.4,
}


def set_font(run, size=BODY_SIZE, bold=False, italic=False, color=None):
    run.font.name = FONT
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    # Ensure East-Asian font mapping uses the same face.
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT)
    rfonts.set(qn("w:hAnsi"), FONT)
    rfonts.set(qn("w:cs"), FONT)


TOKEN = re.compile(r"(\*\*.+?\*\*|\*.+?\*)")


def add_inline(paragraph, text, size=BODY_SIZE, bold=False, italic=False):
    """Render **bold** and *italic* spans into a paragraph."""
    for part in TOKEN.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            set_font(paragraph.add_run(part[2:-2]), size=size, bold=True)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            set_font(paragraph.add_run(part[1:-1]), size=size, italic=True)
        else:
            set_font(paragraph.add_run(part), size=size, bold=bold, italic=italic)


def add_page_number_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    set_font(run, size=Pt(10))
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def body_paragraph(doc, text, **kw):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    add_inline(p, text, **kw)
    return p


def heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    add_inline(p, text, size=Pt(13) if level == 1 else Pt(11), bold=True)
    return p


def add_markdown_table(doc, rows):
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    table.autofit = True
    for ri, row in enumerate(rows):
        for ci in range(ncols):
            cell = table.cell(ri, ci)
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(1)
            para.paragraph_format.space_before = Pt(1)
            add_inline(para, row[ci] if ci < len(row) else "", size=Pt(9.5), bold=(ri == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_figure(doc, image_path, caption):
    img = ROOT / image_path
    if not img.exists():
        body_paragraph(doc, f"[Missing figure: {image_path}]")
        return
    width = FIGURE_WIDTHS.get(img.name, 6.0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(img), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cap.paragraph_format.space_after = Pt(10)
    cap.paragraph_format.first_line_indent = Inches(-0.25)
    cap.paragraph_format.left_indent = Inches(0.25)
    add_inline(cap, caption, size=CAPTION_SIZE)
    cap.runs[0].font.bold = True if caption.startswith("Figure") else False


FIGURE_LINE = re.compile(r"^(?P<cap>\*\*Figure \d+\..*?)\s*\(`?(?P<path>output/figures/[^)`]+)`?\)\s*\*?\s*$")


def main():
    text = MD.read_text(encoding="utf-8")

    doc = Document()

    # Base style.
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = BODY_SIZE
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.0

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    add_page_number_footer(section)

    lines = text.splitlines()

    # ---- Title page: everything up to the first horizontal rule ----
    title_page = []
    idx = 0
    for i, ln in enumerate(lines):
        if ln.strip() == "---":
            idx = i + 1
            break
        title_page.append(ln)

    main_title = None
    for ln in title_page:
        if ln.startswith("# "):
            main_title = ln[2:].strip()
    if main_title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(120)
        p.paragraph_format.space_after = Pt(30)
        p.paragraph_format.line_spacing = 1.15
        add_inline(p, main_title, size=Pt(18), bold=True)

    for ln in title_page:
        if not ln.strip() or ln.startswith("# "):
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        add_inline(p, ln, size=Pt(12))
    # Page break after title page.
    doc.add_page_break()

    # ---- Body ----
    figures_section = False
    in_references = False
    i = idx
    while i < len(lines):
        ln = lines[i]
        stripped = ln.strip()
        if not stripped:
            i += 1
            continue
        if stripped == "---":
            i += 1
            continue
        if stripped.startswith("## "):
            text_h = stripped[3:].strip()
            if text_h.lower().startswith("figures"):
                figures_section = True
                doc.add_page_break()
            in_references = text_h.lower().startswith("references")
            heading(doc, text_h, 1)
            i += 1
            continue
        if stripped.startswith("### "):
            heading(doc, stripped[4:].strip(), 2)
            i += 1
            continue

        m = FIGURE_LINE.match(stripped)
        if m:
            cap = m.group("cap").strip().rstrip("*").strip()
            add_figure(doc, m.group("path").strip(), cap)
            i += 1
            continue

        if stripped.startswith("```"):
            i += 1
            eq_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                if lines[i].strip():
                    eq_lines.append(lines[i].strip())
                i += 1
            i += 1  # skip closing fence
            for el in eq_lines:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(2)
                add_inline(p, el, size=Pt(10.5), italic=True)
            continue

        if stripped.startswith("|"):
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r"-{2,}", c) for c in cells):
                    table_rows.append(cells)
                i += 1
            if table_rows:
                add_markdown_table(doc, table_rows)
            continue

        # Paragraph: collect until blank line.
        para_lines = []
        while i < len(lines) and lines[i].strip():
            para_lines.append(lines[i])
            i += 1
        para_text = " ".join(pl.strip() for pl in para_lines)
        if not para_text:
            continue

        if in_references and re.match(r"^\d+\.", para_text):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Inches(0.30)
            p.paragraph_format.first_line_indent = Inches(-0.30)
            add_inline(p, para_text)
        else:
            body_paragraph(doc, para_text)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
